# -*- coding: utf-8 -*-
import os
import re
import json
import base64
import logging
import traceback
import requests
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


__all__ = ['WriteDocxInference']

class WriteDocxInference(object):

    def __init__(self):

        with open("config.json", encoding="utf-8") as req_source:
            source = json.loads(req_source.read())
            # 报告写作prompt
            self.feasibility_studies_written_config = source["feasibility_studies_written_config"]

        self.candidate_placeholder = ["{{candidate}}"]

    def get_docx_pluging(self, config, template_id):
        """获取写入docx所需要的模板占位符"""
    
        template_path = config["template_style"][template_id]["template_path"]
        placeholders = config["template_style"][template_id]["placeholders"]
        type_id = config["type_id"]
        return template_path, placeholders, type_id

    def placeholder_text_to_docx(self, new_data, template_path, placeholders=None, font_size=12,equel=True):
        """
            将占位符的对应数据写入docx
            :paras new_data: 每个占位符对应的数据
            :paras template_path: 模板路径
            :param equel: 判断new_data模板中的占位符和对应的输出文本是否相等
        """
        logging.info("placeholders is: {}, template_path is: {}".format(placeholders, template_path))
        template = Document(template_path)
        if equel:
            template.styles["Normal"].font.name = u"仿宋"
            template.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
            # template.styles["Normal"].font.size = Pt(font_size) 
            for placeholder in placeholders:
                for paragraph in template.tables[0].rows:
                    if paragraph.cells[0].text.strip() == "{{" + placeholder + "}}":
                        text = new_data[placeholder].replace(placeholder, "")
                        paragraph.cells[0].text = paragraph.cells[0].text.replace("{{" + placeholder + "}}", "")  # 将模板中占位符先替换为空
                        # 添加docx对象
                        para = template.add_paragraph()
                        # 往里面写数据, 之所以先加一个\n. 目前暂时发现如果不加\n写入的材料会自动加空格
                        p = paragraph.cells[0].paragraphs[0].add_run(text + "\n")
                        p.font.size = Pt(font_size)
        else:
            template.styles["Normal"].font.name = u"仿宋"
            template.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
            print('----------new_data---------', new_data)
            candidate_placeholder = ["{{" + "{}".format(list(new_data.keys())[0]) + "}}"]
            print('---------candidate_placeholder-----------', candidate_placeholder)
            # for placeholder in self.candidate_placeholder:
            for placeholder in candidate_placeholder:
                for paragraph in template.tables[3].rows:
                    if paragraph.cells[0].text.strip() == placeholder:
                        paragraph.text = paragraph.text.replace(placeholder, "")
                        para = template.add_paragraph()
                        words = "".join([value for value in list(new_data.values())])
                        p = paragraph.add_run(words + "\n")
                        p.font.size = Pt(font_size)

        return template

    def write_to_template(self, fill_template_dict):
        """
            写入模板, 并发送后台存储, 因为要保证模型结果一定写入到文件, 现在不能保证百分百是按照模板的格式进行输出, 所以需要如果没有按照模板进行输出则不填充, 直接将模型的结果写入
            :paras fill_template_dict: 流式解码所有数据
            :paras file_name: 后端存储的文件名
            :param template_id: 模板id
        """
        try:
            template_path = None
            placeholders = None
            type_id = None
            """
                type_id 是前端和后端规定, 我们不需要做任何修改, 只是做一次转发
                template_path, placeholders: 模板路径和占位符(就是我们需要填充模板在什么位置的标识符)
            """
            template_path, placeholders, type_id = self.get_docx_pluging(self.feasibility_studies_written_config, '0')

            new_data = {}
            for key, value in fill_template_dict.items():
                # 过滤一些模型生成的前缀
                value = re.sub(r"^(研究报告的目的和意义|目的和意义|报告目的和意义|项目研究的背景|研究目标及主要研究内容).*?(\n)", "", value.strip()).lstrip()  # 去掉后无意义的前缀后左边还有一个\n
                # if key in ["二、主要进展", "三、下周重点工作"]:
                #     value = "    " + re.sub(r"\n", "\n    ", value)
                # else:
                #     value = "    " + re.sub(r"\n\b", "\n    ", value)
                # elif r"月报" in file_name:
                #     value = re.sub(r"^(本月|下月).*?(进展：|计划：)", "", value.strip())
                #     value = "    " + re.sub(r"\n\b", "\n    ", value.strip())  # 在\n后面加上缩进

                # elif r"工作汇报" in file_name:
                #     value = "    " + re.sub(r"\n\b", "\n    ", value.strip())  # 在\n后面加上缩进
                fill_template_dict[key] = value

            new_data = fill_template_dict

            logging.info("write docx result is: {}".format(new_data))
            # 说明模型没有按照预定的格式进行输出, 例如模板有3个占位符,但是输出的结构只有一份,有两种情况
            if len(placeholders) != len(new_data):
                """
                    step1: 情况一: 
                            1. 本来就是大模型一次性生成,但是占位符却有多个,这个时候需要我们对生成的结构进行切分成和占位符一样的数量,如果能切分就刚好进行回填
                            2. 切分出来的结果和占位符不一样那同样一次性写入,保证不报错或者写入不了(保证有输出)
                    step2: 情况二: 模型因为某种中间处理出了问题,虽然是要和占位符保持一致的数量,但是在处理过程中导致和占位符不相等,那这种直接合并一次性写入, 保证不报错或者写入不了(保证有输出)
                """
                post_new_data = {}  # 不相等的时候后处理字典
                total_words = "".join([sententce for sententce in list(new_data.values())])
                for i, placeholder in enumerate(placeholders):
                    if i != len(placeholders) - 1:
                        content_reg = re.search(rf"(?<={re.escape(placeholder)})(.|\n)*(?={re.escape(placeholders[i+1])})", total_words)
                    else:
                        content_reg = re.search(rf"(?<={re.escape(placeholder)})(.|\n)*", total_words)

                    if content_reg:
                        post_new_data[placeholder] = content_reg.group()

                if not post_new_data:
                    post_new_data = new_data
                # 重新判断切分后的字典是否和占位符一致
                if len(placeholders) != len(post_new_data):
                    template_path = os.path.join(os.path.dirname(template_path), "condidate.docx")
                    template = self.placeholder_text_to_docx(post_new_data, template_path, equel=False)
                else:
                    template = self.placeholder_text_to_docx(post_new_data, template_path, placeholders=placeholders)
            else:
                template = self.placeholder_text_to_docx(new_data, template_path, placeholders=placeholders)

            template.save("公司科技项目可研报告简化.docx")
            # 创建一个BytesIO对象
            # stream = BytesIO()
            # template.save(stream)
            # file_bytes = stream.getvalue()
            logging.info("Writing template successfully")

            # 发送给java后台服务
            # self.send_file_to_backend(file_bytes, file_name, nw_id, type_id, token_id)

        except Exception as err:
            logging.error(traceback.format_exc())
            raise "Faild to writting docx."  

    def send_file_to_backend(self, file_bytes, file_name, nw_id, type_id, token_id):
        """"""
        try:
            # http://192.168.210.240:18081/fileUpload/saveResultFile
            base64_data = base64.b64encode(file_bytes).decode('utf-8')
            body = {
                "template_id": nw_id,
                "file_name": file_name,
                "base64_data": base64_data,
                "type": type_id
            }
            with open("./3.docx", "wb") as f:
                f.write(file_bytes)
            logging.info("backend request parameter template_id is: {}, file_name is: {}, type_id is: {}".format(nw_id, file_name, type_id))
            headers = {"Content-type": "application/json", "X-Access-Token": token_id}
            response = requests.post(url=self.weekly_written_config["backend_url"], json=body, headers=headers)
            result = json.loads(response.text)
            logging.info("minio response result is: {}".format(result))
            if "id" in result:
                logging.info("Sending the file to the backend server was successful, the url is: {}.".format(self.weekly_written_config["backend_url"]))
            else:
                logging.error("Failed to send files to the backend server, the url is: {}.".format(self.weekly_written_config["backend_url"]))
            return 
        except Exception as err:
            logging.error(traceback.format_exc())
            raise "Failed to send data"


if __name__ == "__main__":
    template = Document("C:\\Users\\zjt78\\Documents\\python_test\\templates\\公司科技项目可研报告.docx")
    print(123)